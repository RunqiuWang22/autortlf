import streamlit as st
import os
import glob
import subprocess
from google import genai
import docx
import pandas as pd

# Define the example YAML
EXAMPLE_YAML = """
type: Table
rfunction: baseline0char
table_id: "Table 12.1.1.x"
rename_output: gemini_baseline_table
title: "Baseline Characteristics"
subtitle: "(GLOBAL.population.ITT.title)"
population_from: "ADSL"
population_filter: GLOBAL.population.ITT.filter_expression
treatment_var: GLOBAL.treatment_config.treatment_random_var
treatment_code: GLOBAL.treatment_config.treatment_random_code_var
display_options:
   display_n: false
   display_mean: true
   display_median: true
   display_range: true
   display_IQR: true
   display_total_column: GLOBAL.formatting.display_options.display_total_column
   display_only_total_column: false
   create_output_dataset: GLOBAL.formatting.display_options.create_output_dataset
output_format:
  output_data_format: GLOBAL.output_format.output_data_format
  output_tlf_format: GLOBAL.output_format.output_tlf_format
decimals:
  continuous: GLOBAL.formatting.decimals.continuous
  percent: GLOBAL.formatting.decimals.percent 
footnotes:
  standard: 
    - GLOBAL.titles.footnotes.standard
  data_cutoff:
    - GLOBAL.titles.footnotes.data_cutoff
data_source_text: "GLOBAL.titles.footnotes.data_source_text adam-adsl]"
variables:
  - name: "Age (years)"
    source_var: "AGE"
    type: "continuous"
  - name: "Sex"
    source_var: "SEX"
    type: "categorical"
    levels: ["M", "F"]
    label_overrides:
      M: "Male"
      F: "Female"
  - name: "Race"
    source_var: "RACE"
    type: "categorical"
"""

EXAMPLE2_YAML_AE = """
type: Table
rfunction: ae0specific
table_id: "Table 14.3.1.1.X"
rename_output: gemini_ae_specific_table
title: "Adverse Events"
subtitle1: "(Incidence >=5% in One or More Treatment Groups)"
subtitle2: "(GLOBAL.population.SAFETY.title)"
population_from: "ADSL"
observation_from: "ADAE"
population_filter: GLOBAL.population.SAFETY.filter_expression
observation_filter: GLOBAL.safety_analysis_param.observation_filter
treatment_var: GLOBAL.treatment_config.treatment_actual_var
display_options:  
   display_total_column: GLOBAL.formatting.display_options.display_total_column
output_format:
  output_data_format: GLOBAL.output_format.output_data_format
  output_tlf_format: GLOBAL.output_format.output_tlf_format 
decimals:
  continuous: GLOBAL.formatting.decimals.continuous
  percent: GLOBAL.formatting.decimals.percent
footnotes:
  standard: 
    - GLOBAL.titles.footnotes.standard
  data_cutoff:
    - GLOBAL.titles.footnotes.data_cutoff
ae_parameters:
  min_subjects_threshold: 5
  ae_term_var: "AEDECOD"
  group_by_var: "AESOC"
  group_display_name: "System Organ Class"
  to_proper_case: true
  sort_options:
    sort_by: "frequency"
    sort_column: "Xanomeline High Dose"
    sort_order: "desc"
    group_sort: "alphabetical"
    within_group_sort: "frequency"
  display_grouping_headers: true
"""

st.set_page_config(page_title="AutoRTLF Table Generator", layout="centered", page_icon="📊")

st.title("📊 AutoRTLF Table Generator")
st.markdown("Easily generate fully formatted RTF tables from table shells using GenAI. No coding required.")

# 1. API Key Input
with st.sidebar:
    st.header("⚙️ Configuration")
    api_key = st.text_input("Gemini API Key", type="password", help="Get a free key at https://aistudio.google.com/app/apikey")
    st.markdown("---")
    st.markdown("*Your API key is not permanently stored. It is only used for the current session.*")

# 2. Main Workflow Tabs
tab1, tab2 = st.tabs(["1. Generate Table Shell from SAP", "2. Generate RTF from Table Shell"])

with tab1:
    st.subheader("Extract Table Shell from SAP Document")
    st.markdown("Upload a Statistical Analysis Plan (SAP) excerpt to automatically design a Markdown Table Shell.")
    sap_file = st.file_uploader("Upload SAP (Word, Text)", type=["text", "txt", "docx"], key="sap_uploader")
    
with tab2:
    st.subheader("Generate RTF Configuration & Table")
    st.markdown("Upload an existing Table Shell to generate the YAML config and the final RTF table.")
    uploaded_file = st.file_uploader("Upload Table Shell (Word, Text, Markdown, CSV)", type=["text", "txt", "md", "docx", "csv"], key="shell_uploader")
    output_yaml_name = st.text_input("Output Configuration Filename", "gemini_baseline_table.yaml")

def extract_text(file):
    if file.name.endswith(".txt") or file.name.endswith(".text") or file.name.endswith(".md"):
        return file.getvalue().decode("utf-8")
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                text += "\n" + " | ".join(row_data)
        return text
    elif file.name.endswith(".csv"):
        df = pd.read_csv(file)
        return df.to_string()
    return ""

def generate_yaml(api_key, table_request, output_filename):
    client = genai.Client(api_key=api_key)
    
    system_instruction = "You are an expert clinical statistical programmer. You only output valid, syntactically correct YAML configurations for the AutoRTLF framework. Do NOT include markdown blocks (```yaml), conversational text, or any explanations. Output ONLY raw YAML text. Make sure to set rename_output to exactly match the requested file output name (without the .yaml extension)."
    prompt = f"Here is the strict example YAML format for 'Baseline Characteristics' (rfunction: baseline0char):\n{EXAMPLE_YAML}\n\nHere is the strict example YAML format for an 'Adverse Event' or 'AE' Table (rfunction: ae0specific):\n{EXAMPLE2_YAML_AE}\n\nPlease generate a new YAML configuration for this request based on the uploaded table shell content. First, carefully read the request and decide if it is a Baseline Characteristics table, an Adverse Event table, or something else, and use the appropriate `rfunction` and overall YAML structure. The YAML MUST match the overall structure exactly. CRITICAL: YOU MUST INCLUDE THE `variables:` OR `ae_parameters:` BLOCK AS SHOWN IN THE EXAMPLES. DO NOT OMIT IT. Inside this block, you must map out the specific variables explicitly present in the uploaded document (e.g., if you see Age and Sex in the document, create the config entries for them). Read the document carefully and extract the requested elements. Use GLOBAL parameters for generic configs where applicable. Set rename_output to '{output_filename.replace('.yaml','')}' \n\n{table_request}"
    
    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt,
        config=genai.types.GenerateContentConfig(
            system_instruction=system_instruction,
            temperature=0.1
        )
    )
    
    generated_yaml = response.text.strip()
    
    # Strip markdown formatting if AI includes it
    if generated_yaml.startswith("```yaml"):
        generated_yaml = generated_yaml[7:]
    elif generated_yaml.startswith("```"):
        generated_yaml = generated_yaml[3:]
    if generated_yaml.endswith("```"):
        generated_yaml = generated_yaml[:-3]
    generated_yaml = generated_yaml.strip()
    
    output_path = os.path.join("pganalysis", "metadata", output_filename)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    with open(output_path, "w") as f:
        f.write(generated_yaml)
        
    return output_path

def get_rtf_by_yaml(yaml_filename):
    # AutoRTLF typically outputs the rtf file with the same base name as the yaml file
    # e.g., baseline0char0itt.yaml -> baseline0char0itt.rtf
    base_name = yaml_filename.replace(".yaml", "")
    
    # We will search the outtable directory for any RTF file matching this base name
    # We use a wildcard for the subfolder because AutoRTLF uses dynamic subfolders (e.g. X99-ia01)
    search_pattern = f"outtable/**/{base_name}.rtf"
    list_of_files = glob.glob(search_pattern, recursive=True)
    
    if list_of_files:
        return list_of_files[0]
    return None

def generate_table_shell(api_key, sap_text):
    client = genai.Client(api_key=api_key)
    
    system_instruction = "You are an expert clinical statistical programmer. Your task is to extract the details from the uploaded Statistical Analysis Plan (SAP) and generate a clear, Markdown-formatted Table Shell. Do not include introductory conversational text."
    prompt = f"Please read the following Statistical Analysis Plan (SAP) excerpt and design a Table Shell for it. Format the Table Shell as a Markdown table, including the title, column headers, stub (row labels), and expected data representation (e.g., 'xx (xx.x%)' for counts, 'xx.x (xx.x)' for mean(SD)).\n\nSAP Content:\n{sap_text}"
    
    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt,
        config=genai.types.GenerateContentConfig(
            system_instruction=system_instruction,
            temperature=0.2
        )
    )
    return response.text.strip()

st.markdown("---")

with tab1:
    if st.button("Extract Table Shell ✨", type="primary", use_container_width=True, key="btn_shell"):
        if not api_key:
            st.error("❌ Please enter your Gemini API Key in the sidebar first.")
        elif not sap_file:
            st.error("❌ Please upload an SAP document.")
        else:
            with st.status("Analyzing SAP...", expanded=True) as status:
                try:
                    st.write("📄 Extracting text from SAP...")
                    sap_text = extract_text(sap_file)
                    
                    st.write("🤖 Generating Table Shell with Gemini AI...")
                    shell_markdown = generate_table_shell(api_key, sap_text)
                    
                    status.update(label="Table Shell Generation Complete!", state="complete", expanded=False)
                    
                    st.success("🎉 Table Shell generated successfully!")
                    st.markdown("### Generated Table Shell")
                    st.markdown(shell_markdown)
                    
                    st.download_button(
                        label="⬇️ Download Table Shell as Markdown",
                        data=shell_markdown,
                        file_name="generated_table_shell.md",
                        mime="text/markdown"
                    )
                except Exception as e:
                    status.update(label="Generation Failed", state="error", expanded=True)
                    st.error(f"An error occurred: {str(e)}")

with tab2:
    if st.button("Generate Table 🚀", type="primary", use_container_width=True, key="btn_rtf"):
        if not api_key:
            st.error("❌ Please enter your Gemini API Key in the sidebar first.")
        elif not uploaded_file:
            st.error("❌ Please upload a Table Shell document.")
        elif not output_yaml_name.endswith(".yaml"):
            st.error("❌ Output filename must end with '.yaml'")
        else:
            with st.status("Processing Request...", expanded=True) as status:
                try:
                    st.write("📄 Extracting text from uploaded file...")
                    table_request = extract_text(uploaded_file)
                    
                    st.write("🤖 Communicating with Gemini AI...")
                    yaml_path = generate_yaml(api_key, table_request, output_yaml_name)
                    st.write(f"✅ AutoRTLF configuration saved to `{yaml_path}`")
                    
                    st.write("⚙️ Running AutoRTLF Data Engine...")
                    
                    # Command 1: Generate batch script
                    subprocess.run(["Rscript", "generate_batch_commands.R"], check=True, capture_output=True, text=True)
                    
                    # Command 2: Execute all commands
                    st.write("🧮 Generating Final RTF Document...")
                    batch_cmd = "cat batch_commands.txt | grep -v '^#' | grep -v '^$' | while read cmd; do eval \"$cmd\"; done"
                    result = subprocess.run(batch_cmd, shell=True, executable='/bin/bash', capture_output=True, text=True)
                    
                    status.update(label="Table Generation Complete!", state="complete", expanded=False)
                    
                    target_rtf = get_rtf_by_yaml(output_yaml_name)
                    
                    if target_rtf:
                        st.success("🎉 Table Built Successfully!")
                        st.write(f"File created: `{os.path.basename(target_rtf)}`")
                        with open(target_rtf, "rb") as file:
                            st.download_button(
                                label="⬇️ Download Your RTF Table",
                                data=file,
                                file_name=os.path.basename(target_rtf),
                                mime="application/rtf",
                                type="primary",
                                key="btn_download_rtf"
                            )
                    else:
                        st.error("⚠️ No RTF file was generated. Please check the logic of the prompt.")
                        
                except Exception as e:
                    status.update(label="Generation Failed", state="error", expanded=True)
                    st.error(f"An error occurred: {str(e)}")
